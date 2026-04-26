import base64
import os

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.db.models import Count, Avg, Q


def _img_b64(relative_static_path):
    """Return base64-encoded content of a static image file, or '' on error."""
    from django.conf import settings as _s
    full = os.path.join(_s.STATICFILES_DIRS[0] if _s.STATICFILES_DIRS else _s.STATIC_ROOT or '', relative_static_path)
    if not os.path.isfile(full):
        # fallback: look inside each STATICFILES_DIRS entry
        for d in getattr(_s, 'STATICFILES_DIRS', []):
            candidate = os.path.join(d, relative_static_path)
            if os.path.isfile(candidate):
                full = candidate
                break
        else:
            return ''
    with open(full, 'rb') as f:
        return base64.b64encode(f.read()).decode('ascii')
from .models import (
    Competency, CompetencyCategory, EmployeeCompetency, PositionCompetency,
    ProficiencyLevel,
    IndividualDevelopmentPlan, IDPActivity,
    JobDescription, JobAnalysis, JobAnalysisCompetency,
    JobAnalysisEntry, SecondaryDuty, RequiredSkill, ToolEquipment,
    DOTrCompetency, DOTrCompetencyIndicator, DOTrCompetencyType, DOTrOfficeMandate,
)
from apps.accounts.models import Division
from apps.accounts.decorators import role_required
from django.contrib.auth import get_user_model
User = get_user_model()


@login_required
def competency_list(request):
    competencies = Competency.objects.filter(is_active=True).order_by('category', 'name')
    return render(request, 'competencies/competency_list.html', {'competencies': competencies})


@login_required
@role_required(['admin', 'hr'])
def competency_create(request):
    if request.method == 'POST':
        name = request.POST.get('name')
        code = request.POST.get('code')
        category = request.POST.get('category')
        description = request.POST.get('description')
        behavioral_indicators = request.POST.get('behavioral_indicators', '')
        Competency.objects.create(
            name=name, code=code, category=category,
            description=description, behavioral_indicators=behavioral_indicators
        )
        messages.success(request, 'Competency created successfully.')
        return redirect('competency_list')
    from .models import CompetencyCategory
    return render(request, 'competencies/competency_form.html', {
        'title': 'Add Competency',
        'categories': CompetencyCategory.choices
    })


@login_required
@role_required(['admin', 'hr'])
def competency_edit(request, pk):
    competency = get_object_or_404(Competency, pk=pk)
    if request.method == 'POST':
        competency.name = request.POST.get('name')
        competency.code = request.POST.get('code')
        competency.category = request.POST.get('category')
        competency.description = request.POST.get('description')
        competency.behavioral_indicators = request.POST.get('behavioral_indicators', '')
        competency.save()
        messages.success(request, 'Competency updated.')
        return redirect('competency_list')
    from .models import CompetencyCategory
    return render(request, 'competencies/competency_form.html', {
        'title': 'Edit Competency',
        'competency': competency,
        'categories': CompetencyCategory.choices
    })


@login_required
def my_competencies(request):
    user = request.user
    employee_comps = (
        EmployeeCompetency.objects
        .filter(user=user)
        .select_related('competency', 'jaf_entry')
        .order_by('source', 'competency__category', 'competency__name')
    )
    all_competencies = Competency.objects.filter(is_active=True)
    from apps.trainings.models import TrainingProgram
    recommended = TrainingProgram.objects.filter(status='published').order_by('?')[:5]
    return render(request, 'competencies/my_competencies.html', {
        'employee_comps': employee_comps,
        'all_competencies': all_competencies,
        'recommended': recommended,
    })


@login_required
@role_required(['admin', 'hr', 'supervisor'])
def employee_competency_view(request, user_pk):
    profile_user = get_object_or_404(User, pk=user_pk)
    employee_comps = EmployeeCompetency.objects.filter(user=profile_user).select_related('competency')
    position_comps = PositionCompetency.objects.filter(position=profile_user.position).select_related('competency')
    return render(request, 'competencies/employee_competencies.html', {
        'profile_user': profile_user,
        'employee_comps': employee_comps,
        'position_comps': position_comps,
    })


@login_required
@role_required(['admin', 'hr', 'supervisor'])
def update_employee_competency(request, user_pk):
    profile_user = get_object_or_404(User, pk=user_pk)
    if request.method == 'POST':
        competency_id = request.POST.get('competency')
        current_level = request.POST.get('current_level')
        target_level = request.POST.get('target_level')
        notes = request.POST.get('notes', '')
        import datetime
        ec, created = EmployeeCompetency.objects.update_or_create(
            user=profile_user,
            competency_id=competency_id,
            defaults={
                'current_level': current_level,
                'target_level': target_level,
                'notes': notes,
                'assessed_by': request.user,
                'assessment_date': datetime.date.today()
            }
        )
        messages.success(request, 'Competency rating updated.')
    return redirect('employee_competency_view', user_pk=user_pk)


@login_required
def gap_analysis(request):
    """Competency gap analysis for current user or all users (HR/Admin)."""
    if request.user.role in ['admin', 'hr']:
        users = User.objects.filter(is_active=True).prefetch_related('competencies__competency')
        gaps = []
        for user in users:
            for ec in user.competencies.all():
                gap = ec.get_gap()
                if gap > 0:
                    gaps.append({'user': user, 'competency': ec.competency, 'gap': gap, 'current': ec.current_level, 'target': ec.target_level})
        return render(request, 'competencies/gap_analysis.html', {'gaps': gaps, 'all_users': True})
    else:
        employee_comps = EmployeeCompetency.objects.filter(user=request.user).select_related('competency')
        gaps = [ec for ec in employee_comps if ec.get_gap() > 0]
        return render(request, 'competencies/gap_analysis.html', {'gaps': gaps, 'all_users': False})


@login_required
def idp_list(request):
    if request.user.role in ['admin', 'hr']:
        idps = IndividualDevelopmentPlan.objects.select_related('user').all()
    elif request.user.role == 'supervisor':
        subordinate_ids = request.user.subordinates.values_list('id', flat=True)
        idps = IndividualDevelopmentPlan.objects.filter(user_id__in=subordinate_ids).select_related('user')
    else:
        idps = IndividualDevelopmentPlan.objects.filter(user=request.user)
    return render(request, 'competencies/idp_list.html', {'idps': idps})


@login_required
def idp_create(request):
    import datetime
    if request.method == 'POST':
        year = request.POST.get('year', datetime.date.today().year)
        career_objective = request.POST.get('career_objective', '')
        existing = IndividualDevelopmentPlan.objects.filter(user=request.user, year=year).first()
        if existing:
            messages.warning(request, f'You already have an IDP for {year}.')
            return redirect('idp_detail', pk=existing.pk)
        idp = IndividualDevelopmentPlan.objects.create(
            user=request.user, year=year, career_objective=career_objective
        )
        messages.success(request, f'IDP for {year} created.')
        return redirect('idp_detail', pk=idp.pk)
    return render(request, 'competencies/idp_form.html', {'title': 'Create IDP', 'year': datetime.date.today().year})


@login_required
def idp_detail(request, pk):
    idp = get_object_or_404(IndividualDevelopmentPlan, pk=pk)
    if idp.user != request.user and request.user.role not in ['admin', 'hr', 'supervisor']:
        messages.error(request, 'Access denied.')
        return redirect('idp_list')
    if request.method == 'POST' and 'submit_idp' in request.POST:
        if idp.user == request.user and idp.status == 'draft':
            idp.status = 'submitted'
            idp.save()
            messages.success(request, 'IDP submitted for approval.')
        return redirect('idp_detail', pk=pk)
    activities = idp.activities.select_related('competency').all()
    competencies = Competency.objects.filter(is_active=True)
    return render(request, 'competencies/idp_detail.html', {
        'idp': idp, 'activities': activities, 'competencies': competencies
    })


@login_required
def idp_add_activity(request, idp_pk):
    idp = get_object_or_404(IndividualDevelopmentPlan, pk=idp_pk, user=request.user)
    if idp.status == 'completed':
        messages.error(request, 'Cannot add activities to a completed IDP.')
        return redirect('idp_detail', pk=idp_pk)
    if request.method == 'POST':
        IDPActivity.objects.create(
            idp=idp,
            competency_id=request.POST.get('competency'),
            learning_intervention=request.POST.get('learning_intervention'),
            timeline=request.POST.get('timeline'),
            success_indicator=request.POST.get('success_indicator', '')
        )
        messages.success(request, 'Activity added to IDP.')
    return redirect('idp_detail', pk=idp_pk)


@login_required
def idp_toggle_activity(request, activity_pk):
    if request.method != 'POST':
        return redirect('idp_list')
    activity = get_object_or_404(IDPActivity, pk=activity_pk)
    idp = activity.idp
    # Employee can mark their own activities; supervisors/hr/admin can mark any
    if idp.user != request.user and request.user.role not in ['admin', 'hr', 'supervisor']:
        messages.error(request, 'Access denied.')
        return redirect('idp_list')
    activity.is_completed = not activity.is_completed
    activity.save()
    return redirect('idp_detail', pk=idp.pk)


@login_required
@role_required(['admin', 'hr', 'supervisor'])
def idp_approve(request, pk):
    import datetime
    idp = get_object_or_404(IndividualDevelopmentPlan, pk=pk)
    idp.status = 'approved'
    idp.approved_by = request.user
    idp.approval_date = datetime.date.today()
    idp.save()
    messages.success(request, f"IDP for {idp.user.get_full_name()} approved.")
    return redirect('idp_list')


# ── Job Description ───────────────────────────────────────────────────────────

@login_required
@role_required(['admin', 'hr'])
def job_description_list(request):
    job_descriptions = JobDescription.objects.filter(is_active=True).select_related('division', 'created_by')
    return render(request, 'competencies/job_description_list.html', {
        'job_descriptions': job_descriptions,
    })


@login_required
@role_required(['admin', 'hr'])
def job_description_create(request):
    from apps.accounts.models import Division
    if request.method == 'POST':
        jd = JobDescription.objects.create(
            position_title=request.POST.get('position_title'),
            division_id=request.POST.get('division') or None,
            salary_grade=request.POST.get('salary_grade') or None,
            employment_status=request.POST.get('employment_status', ''),
            office_unit=request.POST.get('office_unit', ''),
            immediate_supervisor=request.POST.get('immediate_supervisor', ''),
            duties_and_responsibilities=request.POST.get('duties_and_responsibilities', ''),
            performance_standards=request.POST.get('performance_standards', ''),
            education=request.POST.get('education', ''),
            training=request.POST.get('training', ''),
            experience=request.POST.get('experience', ''),
            eligibility=request.POST.get('eligibility', ''),
            knowledge_skills_abilities=request.POST.get('knowledge_skills_abilities', ''),
            created_by=request.user,
        )
        messages.success(request, f'Job description for "{jd.position_title}" created.')
        return redirect('job_description_detail', pk=jd.pk)
    divisions = Division.objects.select_related('head').all()
    division_heads = {
        str(d.pk): d.head.get_full_name() if d.head else ''
        for d in divisions
    }
    return render(request, 'competencies/job_description_form.html', {
        'title': 'Create Job Description',
        'divisions': divisions,
        'division_heads': division_heads,
    })


@login_required
@role_required(['admin', 'hr'])
def job_description_detail(request, pk):
    jd = get_object_or_404(JobDescription, pk=pk)
    analyses = jd.analyses.select_related('analyzed_by').order_by('-created_at')
    return render(request, 'competencies/job_description_detail.html', {
        'jd': jd,
        'analyses': analyses,
    })


@login_required
@role_required(['admin', 'hr'])
def job_description_edit(request, pk):
    from apps.accounts.models import Division
    jd = get_object_or_404(JobDescription, pk=pk)
    if request.method == 'POST':
        jd.position_title = request.POST.get('position_title')
        jd.division_id = request.POST.get('division') or None
        jd.salary_grade = request.POST.get('salary_grade') or None
        jd.employment_status = request.POST.get('employment_status', '')
        jd.office_unit = request.POST.get('office_unit', '')
        jd.immediate_supervisor = request.POST.get('immediate_supervisor', '')
        jd.duties_and_responsibilities = request.POST.get('duties_and_responsibilities', '')
        jd.performance_standards = request.POST.get('performance_standards', '')
        jd.education = request.POST.get('education', '')
        jd.training = request.POST.get('training', '')
        jd.experience = request.POST.get('experience', '')
        jd.eligibility = request.POST.get('eligibility', '')
        jd.knowledge_skills_abilities = request.POST.get('knowledge_skills_abilities', '')
        jd.save()
        messages.success(request, 'Job description updated.')
        return redirect('job_description_detail', pk=jd.pk)
    divisions = Division.objects.select_related('head').all()
    division_heads = {
        str(d.pk): d.head.get_full_name() if d.head else ''
        for d in divisions
    }
    return render(request, 'competencies/job_description_form.html', {
        'title': 'Edit Job Description',
        'jd': jd,
        'divisions': divisions,
        'division_heads': division_heads,
    })


# ── Job Analysis ──────────────────────────────────────────────────────────────

@login_required
@role_required(['admin', 'hr'])
def job_analysis_create(request, jd_pk):
    import datetime
    jd = get_object_or_404(JobDescription, pk=jd_pk)
    if request.method == 'POST':
        analysis = JobAnalysis.objects.create(
            job_description=jd,
            analyzed_by=request.user,
            analysis_date=datetime.date.today(),
            summary=request.POST.get('summary', ''),
        )
        messages.success(request, 'Job analysis started. Add competency requirements below.')
        return redirect('job_analysis_detail', pk=analysis.pk)
    return render(request, 'competencies/job_analysis_form.html', {'jd': jd})


@login_required
@role_required(['admin', 'hr'])
def job_analysis_detail(request, pk):
    analysis = get_object_or_404(JobAnalysis, pk=pk)
    from .models import CompetencyCategory, ProficiencyLevel
    mapped_ids = analysis.competency_mappings.values_list('competency_id', flat=True)
    available_competencies = Competency.objects.filter(is_active=True).exclude(pk__in=mapped_ids)
    return render(request, 'competencies/job_analysis_detail.html', {
        'analysis': analysis,
        'mappings': analysis.competency_mappings.select_related('competency').order_by(
            'competency__category', 'competency__name'
        ),
        'available_competencies': available_competencies,
        'proficiency_levels': ProficiencyLevel.choices,
        'categories': CompetencyCategory.choices,
    })


@login_required
@role_required(['admin', 'hr'])
def job_analysis_add_competency(request, pk):
    if request.method != 'POST':
        return redirect('job_analysis_detail', pk=pk)
    analysis = get_object_or_404(JobAnalysis, pk=pk)
    if analysis.status == JobAnalysis.Status.APPLIED:
        messages.error(request, 'Cannot modify a job analysis that has already been applied.')
        return redirect('job_analysis_detail', pk=pk)
    competency_id = request.POST.get('competency')
    required_level = request.POST.get('required_level')
    justification = request.POST.get('justification', '')
    if competency_id and required_level:
        _, created = JobAnalysisCompetency.objects.update_or_create(
            job_analysis=analysis,
            competency_id=competency_id,
            defaults={'required_level': required_level, 'justification': justification},
        )
        competency = get_object_or_404(Competency, pk=competency_id)
        messages.success(request, f'"{competency.name}" added to job analysis.')
    return redirect('job_analysis_detail', pk=pk)


@login_required
@role_required(['admin', 'hr'])
def job_analysis_remove_competency(request, pk, mapping_pk):
    if request.method != 'POST':
        return redirect('job_analysis_detail', pk=pk)
    analysis = get_object_or_404(JobAnalysis, pk=pk)
    if analysis.status == JobAnalysis.Status.APPLIED:
        messages.error(request, 'Cannot modify a job analysis that has already been applied.')
        return redirect('job_analysis_detail', pk=pk)
    mapping = get_object_or_404(JobAnalysisCompetency, pk=mapping_pk, job_analysis=analysis)
    name = mapping.competency.name
    mapping.delete()
    messages.success(request, f'"{name}" removed from job analysis.')
    return redirect('job_analysis_detail', pk=pk)


@login_required
@role_required(['admin', 'hr'])
def job_analysis_finalize(request, pk):
    if request.method != 'POST':
        return redirect('job_analysis_detail', pk=pk)
    analysis = get_object_or_404(JobAnalysis, pk=pk)
    if analysis.status == JobAnalysis.Status.DRAFT:
        analysis.status = JobAnalysis.Status.FINALIZED
        analysis.save()
        messages.success(request, 'Job analysis finalized.')
    return redirect('job_analysis_detail', pk=pk)


@login_required
@role_required(['admin', 'hr'])
def job_analysis_apply(request, pk):
    """Apply finalized job analysis competencies to PositionCompetency for this position."""
    if request.method != 'POST':
        return redirect('job_analysis_detail', pk=pk)
    import datetime
    from django.utils import timezone
    analysis = get_object_or_404(JobAnalysis, pk=pk)
    if analysis.status != JobAnalysis.Status.FINALIZED:
        messages.error(request, 'Only finalized analyses can be applied.')
        return redirect('job_analysis_detail', pk=pk)

    position = analysis.job_description.position_title
    created_count = 0
    updated_count = 0
    for mapping in analysis.competency_mappings.select_related('competency'):
        _, created = PositionCompetency.objects.update_or_create(
            position=position,
            competency=mapping.competency,
            defaults={'required_level': mapping.required_level},
        )
        if created:
            created_count += 1
        else:
            updated_count += 1

    analysis.status = JobAnalysis.Status.APPLIED
    analysis.applied_by = request.user
    analysis.applied_at = timezone.now()
    analysis.save()

    messages.success(
        request,
        f'Job analysis applied: {created_count} new competency requirement(s) created, '
        f'{updated_count} updated for position "{position}".'
    )
    return redirect('job_analysis_detail', pk=pk)


# ── DOTr Job Analysis Form (Employee Self-Report) ─────────────────────────────

def _jaf_is_supervisor_of(user, employee):
    """Return True if user is a supervisor in the same division as the employee."""
    if not user.division_id:
        return False
    return employee.division_id == user.division_id


def _jaf_is_reviewer(user, entry):
    """Return True if the user may act as reviewer (supervisor or division head) for this entry."""
    return _jaf_is_supervisor_of(user, entry.employee)


def _jaf_can_view(user, entry):
    """Return True if the user is allowed to view this entry."""
    if user.role in ['admin', 'hr', 'executive']:
        return True
    if entry.employee_id == user.pk:
        return True
    if _jaf_is_reviewer(user, entry):
        return True
    return False


@login_required
def jaf_list(request):
    """List Job Analysis Form entries visible to the current user."""
    user = request.user
    if user.role in ['admin', 'hr', 'executive']:
        entries = JobAnalysisEntry.objects.select_related('employee').all()
    elif user.role == 'supervisor':
        # Only entries from employees in the supervisor's own division
        if user.division_id:
            entries = JobAnalysisEntry.objects.select_related('employee').filter(
                employee__division_id=user.division_id
            )
        else:
            entries = JobAnalysisEntry.objects.filter(employee=user)
    else:
        entries = JobAnalysisEntry.objects.filter(employee=user)
    return render(request, 'competencies/jaf_list.html', {'entries': entries})


@login_required
def jaf_create(request):
    """Create a new Job Analysis Form entry (employee only fills their own)."""
    user = request.user
    if request.method == 'POST':
        import datetime
        entry = JobAnalysisEntry.objects.create(
            employee=user,
            full_name=request.POST.get('full_name', user.get_full_name()),
            position_title=request.POST.get('position_title', user.position),
            office_service_division=request.POST.get('office_service_division', ''),
            section_project_unit=request.POST.get('section_project_unit', ''),
            alternate_position=request.POST.get('alternate_position', ''),
            job_purpose=request.POST.get('job_purpose', ''),
            main_duties=request.POST.get('main_duties', ''),
            challenges_critical_issues=request.POST.get('challenges_critical_issues', ''),
            additional_comments=request.POST.get('additional_comments', ''),
        )
        _save_jaf_employee_sections(request, entry)
        messages.success(request, 'Job Analysis Form saved as draft.')
        return redirect('jaf_detail', pk=entry.pk)

    division_name = str(user.division) if user.division else ''
    return render(request, 'competencies/jaf_form.html', {
        'title': 'New Job Analysis Form',
        'prefill_full_name': user.get_full_name(),
        'prefill_position_title': user.position,
        'prefill_office_service_division': division_name,
    })


@login_required
def jaf_edit(request, pk):
    """Edit employee sections of a draft / rejected Job Analysis Form."""
    entry = get_object_or_404(JobAnalysisEntry, pk=pk)

    is_owner = entry.employee_id == request.user.pk
    is_privileged = request.user.role in ['admin', 'hr']

    if not is_owner and not is_privileged:
        messages.error(request, 'Access denied.')
        return redirect('jaf_list')

    editable_statuses = ['draft', 'rejected']
    if entry.status not in editable_statuses and not is_privileged:
        messages.error(request, 'Only draft or returned forms can be edited.')
        return redirect('jaf_detail', pk=pk)

    if request.method == 'POST':
        entry.full_name = request.POST.get('full_name', entry.full_name)
        entry.position_title = request.POST.get('position_title', entry.position_title)
        entry.office_service_division = request.POST.get('office_service_division', '')
        entry.section_project_unit = request.POST.get('section_project_unit', '')
        entry.alternate_position = request.POST.get('alternate_position', '')
        entry.job_purpose = request.POST.get('job_purpose', '')
        entry.main_duties = request.POST.get('main_duties', '')
        entry.challenges_critical_issues = request.POST.get('challenges_critical_issues', '')
        entry.additional_comments = request.POST.get('additional_comments', '')
        # Reset to draft when a rejected form is re-edited
        if entry.status == 'rejected':
            entry.status = 'draft'
            entry.rejection_comment = ''
        entry.save()
        entry.secondary_duties.all().delete()
        entry.tools_equipment.all().delete()
        _save_jaf_employee_sections(request, entry)
        messages.success(request, 'Job Analysis Form updated.')
        return redirect('jaf_detail', pk=pk)

    secondary_duties = list(entry.secondary_duties.values('task', 'frequency', 'order'))
    tools = list(entry.tools_equipment.values('name', 'order'))
    return render(request, 'competencies/jaf_form.html', {
        'title': 'Edit Job Analysis Form',
        'entry': entry,
        'secondary_duties': secondary_duties,
        'tools': tools,
        'prefill_full_name': '',
        'prefill_position_title': '',
        'prefill_office_service_division': '',
    })


@login_required
def jaf_detail(request, pk):
    """View a Job Analysis Form entry."""
    entry = get_object_or_404(JobAnalysisEntry, pk=pk)
    if not _jaf_can_view(request.user, entry):
        messages.error(request, 'Access denied.')
        return redirect('jaf_list')
    is_reviewer = _jaf_is_reviewer(request.user, entry)
    revision_comments = entry.revision_comments.select_related('commented_by').all()
    return render(request, 'competencies/jaf_detail.html', {
        'entry': entry,
        'secondary_duties': entry.secondary_duties.all(),
        'required_skills': entry.required_skills.select_related('competency').all(),
        'tools': entry.tools_equipment.all(),
        'revision_comments': revision_comments,
        'can_edit': (entry.employee_id == request.user.pk and entry.status in ['draft', 'rejected'])
                    or request.user.role in ['admin', 'hr'],
        'can_submit': entry.employee_id == request.user.pk and entry.status in ['draft', 'rejected'],
        # Supervisor: "Start Review" button (submitted → supervisor_review)
        'can_start_supervisor_review': is_reviewer and entry.status == 'submitted',
        # Supervisor: "Fill Competencies & Forward to HRDD" (submitted or supervisor_review)
        'can_review': (is_reviewer or request.user.role in ['admin', 'hr'])
                      and entry.status in ('submitted', 'supervisor_review'),
        # HRDD: "Start Review" (pending_hrdd → hrdd_review)
        'can_start_hrdd_review': request.user.role in ['hr', 'admin'] and entry.status == 'pending_hrdd',
        # HRDD: Approve / Reject
        'can_approve': request.user.role in ['hr', 'admin'] and entry.status in ('pending_hrdd', 'hrdd_review'),
        'can_reject': request.user.role in ['hr', 'admin'] and entry.status in ('pending_hrdd', 'hrdd_review'),
        'can_store': request.user.role in ['hr', 'admin'],
    })


@login_required
def jaf_submit(request, pk):
    """Employee certifies and submits the form to their supervisor."""
    if request.method != 'POST':
        return redirect('jaf_detail', pk=pk)
    import datetime
    entry = get_object_or_404(JobAnalysisEntry, pk=pk)
    if entry.employee_id != request.user.pk:
        messages.error(request, 'Access denied.')
        return redirect('jaf_list')
    if entry.status in ['draft', 'rejected']:
        entry.status = 'submitted'
        entry.certified_date = datetime.date.today()
        entry.rejection_comment = ''
        entry.save()
        messages.success(request, 'Job Analysis Form submitted. Awaiting supervisor review.')
    return redirect('jaf_detail', pk=pk)


@login_required
def jaf_review(request, pk):
    """Supervisor / Division Head fills Required Competencies and forwards to HRDD."""
    entry = get_object_or_404(JobAnalysisEntry, pk=pk)

    if not (_jaf_is_reviewer(request.user, entry) or request.user.role in ['admin', 'hr']):
        messages.error(request, 'Only the immediate supervisor or division head may review this form.')
        return redirect('jaf_detail', pk=pk)

    if entry.status not in ('submitted', 'supervisor_review'):
        messages.error(request, 'Only submitted forms can be reviewed.')
        return redirect('jaf_detail', pk=pk)

    if request.method == 'POST':
        import datetime
        entry.required_skills.all().delete()
        _save_jaf_skills(request, entry)
        entry.reviewed_by = request.user
        entry.reviewed_date = datetime.date.today()
        entry.status = JobAnalysisEntry.Status.PENDING_HRDD
        entry.save()
        messages.success(request, 'Competencies recorded. Form forwarded to HRDD for review and approval.')
        return redirect('jaf_detail', pk=pk)

    from .models import CompetencyCategory
    comp_groups = [
        {'label': label, 'items': list(Competency.objects.filter(is_active=True, category=cat).order_by('name'))}
        for cat, label in CompetencyCategory.choices
    ]
    comp_groups = [g for g in comp_groups if g['items']]
    existing_skills = list(
        entry.required_skills.select_related('competency').values(
            'skill_name', 'proficiency_level', 'order', 'competency_id'
        )
    )
    return render(request, 'competencies/jaf_review_form.html', {
        'entry': entry,
        'existing_skills': existing_skills,
        'comp_groups': comp_groups,
    })


@login_required
def jaf_approve(request, pk):
    """HRDD approves the form."""
    if request.method != 'POST':
        return redirect('jaf_detail', pk=pk)
    import datetime
    entry = get_object_or_404(JobAnalysisEntry, pk=pk)
    if request.user.role not in ['hr', 'admin']:
        messages.error(request, 'Only HRDD may approve this form.')
        return redirect('jaf_detail', pk=pk)
    if entry.status not in ('pending_hrdd', 'hrdd_review'):
        messages.error(request, 'Only forms submitted to HRDD can be approved.')
        return redirect('jaf_detail', pk=pk)
    entry.status = JobAnalysisEntry.Status.APPROVED
    entry.approved_by = request.user
    entry.approved_date = datetime.date.today()
    entry.save()

    # Sync required competencies → employee's My Competency Profile
    _LEVEL_MAP = {
        'basic':        ProficiencyLevel.BASIC,
        'intermediate': ProficiencyLevel.INTERMEDIATE,
        'advanced':     ProficiencyLevel.ADVANCED,
        'superior':     ProficiencyLevel.EXPERT,
    }
    synced = 0
    for skill in entry.required_skills.select_related('competency').all():
        if not skill.competency_id:
            continue
        target = _LEVEL_MAP.get(skill.proficiency_level, ProficiencyLevel.BASIC)
        emp_comp, created = EmployeeCompetency.objects.get_or_create(
            user=entry.employee,
            competency=skill.competency,
            defaults={
                'current_level':   ProficiencyLevel.BASIC,
                'target_level':    target,
                'assessed_by':     request.user,
                'assessment_date': datetime.date.today(),
                'source':          EmployeeCompetency.Source.JAF,
                'jaf_entry':       entry,
                'notes': (
                    f'Required for position: {entry.position_title}. '
                    f'Set by {request.user.get_full_name()} from approved JAF #{entry.pk}.'
                ),
            }
        )
        if not created:
            emp_comp.target_level    = target
            emp_comp.assessed_by     = request.user
            emp_comp.assessment_date = datetime.date.today()
            emp_comp.source          = EmployeeCompetency.Source.JAF
            emp_comp.jaf_entry       = entry
            emp_comp.notes = (
                f'Required for position: {entry.position_title}. '
                f'Updated by {request.user.get_full_name()} from approved JAF #{entry.pk}.'
            )
            emp_comp.save()
        synced += 1

    messages.success(
        request,
        f'Job Analysis Form for {entry.full_name} approved. '
        f'{synced} required competenc{"y" if synced == 1 else "ies"} added to their competency profile.'
    )
    return redirect('jaf_detail', pk=pk)


@login_required
def jaf_reject(request, pk):
    """HRDD returns the form to the employee for revision."""
    if request.method != 'POST':
        return redirect('jaf_detail', pk=pk)
    entry = get_object_or_404(JobAnalysisEntry, pk=pk)
    if request.user.role not in ['hr', 'admin']:
        messages.error(request, 'Only HRDD may return a form.')
        return redirect('jaf_detail', pk=pk)
    if entry.status not in ('pending_hrdd', 'hrdd_review'):
        messages.error(request, 'Only forms submitted to HRDD can be returned.')
        return redirect('jaf_detail', pk=pk)
    from .models import JAFRevisionComment
    comment = request.POST.get('rejection_comment', '').strip()
    entry.status = JobAnalysisEntry.Status.REJECTED
    entry.rejection_comment = comment
    entry.approved_by = None
    entry.approved_date = None
    entry.save()
    if comment:
        JAFRevisionComment.objects.create(
            entry=entry,
            comment=comment,
            commented_by=request.user,
        )
    messages.warning(request, f'Form returned to {entry.full_name} for revision.')
    return redirect('jaf_detail', pk=pk)


@login_required
def jaf_employee_picker(request):
    """
    Supervisors / Division Heads pick which employee to create a JAF for.
    Admin and HR see all active employees.
    """
    user = request.user
    if user.role in ['admin', 'hr']:
        employees = User.objects.filter(is_active=True).exclude(pk=user.pk).select_related('division')
    elif user.role == 'supervisor':
        if user.division_id:
            employees = User.objects.filter(
                is_active=True,
                division_id=user.division_id,
            ).exclude(pk=user.pk).select_related('division')
        else:
            employees = User.objects.none()
    else:
        messages.error(request, 'Only supervisors, division heads, HR, or admins may create forms for other employees.')
        return redirect('jaf_list')

    return render(request, 'competencies/jaf_employee_picker.html', {'employees': employees})


@login_required
def jaf_create_for_employee(request, employee_pk):
    """
    Supervisor / Division Head creates and fills a complete JAF on behalf of an employee.
    All sections — including Required Competencies — are editable.
    The saved form is set to 'reviewed' status so it goes straight to the Director.
    """
    import datetime
    employee = get_object_or_404(User, pk=employee_pk, is_active=True)
    user = request.user

    # Access check — supervisor can only create for employees in their own division
    if user.role not in ['admin', 'hr'] and not _jaf_is_supervisor_of(user, employee):
        messages.error(request, 'You can only create forms for employees in your division.')
        return redirect('jaf_employee_picker')

    if request.method == 'POST':
        entry = JobAnalysisEntry.objects.create(
            employee=employee,
            full_name=request.POST.get('full_name', employee.get_full_name()),
            position_title=request.POST.get('position_title', employee.position),
            office_service_division=request.POST.get('office_service_division', ''),
            section_project_unit=request.POST.get('section_project_unit', ''),
            alternate_position=request.POST.get('alternate_position', ''),
            job_purpose=request.POST.get('job_purpose', ''),
            main_duties=request.POST.get('main_duties', ''),
            challenges_critical_issues=request.POST.get('challenges_critical_issues', ''),
            additional_comments=request.POST.get('additional_comments', ''),
            # Supervisor-created forms skip straight to HRDD
            status=JobAnalysisEntry.Status.PENDING_HRDD,
            supervisor_created=True,
            certified_date=datetime.date.today(),
            reviewed_by=user,
            reviewed_date=datetime.date.today(),
        )
        _save_jaf_employee_sections(request, entry)
        _save_jaf_skills(request, entry)
        messages.success(
            request,
            f'Job Analysis Form for {employee.get_full_name()} created and forwarded to HRDD for review and approval.'
        )
        return redirect('jaf_detail', pk=entry.pk)

    from .models import CompetencyCategory
    comp_groups = [
        {'label': label, 'items': list(Competency.objects.filter(is_active=True, category=cat).order_by('name'))}
        for cat, label in CompetencyCategory.choices
    ]
    comp_groups = [g for g in comp_groups if g['items']]
    division_name = str(employee.division) if employee.division else ''
    return render(request, 'competencies/jaf_form.html', {
        'title': f'Create JAF for {employee.get_full_name()}',
        'supervisor_mode': True,
        'for_employee': employee,
        'prefill_full_name': employee.get_full_name(),
        'prefill_position_title': employee.position,
        'prefill_office_service_division': division_name,
        'comp_groups': comp_groups,
    })


@login_required
def jaf_supervisor_start_review(request, pk):
    """Supervisor marks the form as under review (submitted → supervisor_review)."""
    if request.method != 'POST':
        return redirect('jaf_detail', pk=pk)
    entry = get_object_or_404(JobAnalysisEntry, pk=pk)
    if not _jaf_is_reviewer(request.user, entry):
        messages.error(request, 'Access denied.')
        return redirect('jaf_detail', pk=pk)
    if entry.status != 'submitted':
        messages.error(request, 'Only submitted forms can be moved to supervisor review.')
        return redirect('jaf_detail', pk=pk)
    entry.status = JobAnalysisEntry.Status.SUPERVISOR_REVIEW
    entry.reviewed_by = request.user
    import datetime
    entry.reviewed_date = datetime.date.today()
    entry.save()
    messages.info(request, 'Form marked as under supervisor review.')
    return redirect('jaf_detail', pk=pk)


@login_required
def jaf_hrdd_start_review(request, pk):
    """HRDD marks the form as under HRDD review (pending_hrdd → hrdd_review)."""
    if request.method != 'POST':
        return redirect('jaf_detail', pk=pk)
    entry = get_object_or_404(JobAnalysisEntry, pk=pk)
    if request.user.role not in ['hr', 'admin']:
        messages.error(request, 'Only HRDD may start review.')
        return redirect('jaf_detail', pk=pk)
    if entry.status != 'pending_hrdd':
        messages.error(request, 'Form must be in "Submitted to HRDD" status to start review.')
        return redirect('jaf_detail', pk=pk)
    entry.status = JobAnalysisEntry.Status.HRDD_REVIEW
    entry.save()
    messages.info(request, 'Form marked as under HRDD review.')
    return redirect('jaf_detail', pk=pk)


@login_required
@role_required(['admin', 'hr'])
def jaf_store_index(request):
    """List all stored JAF HTML files grouped by division folder (Admin/HRDD only)."""
    from django.conf import settings as django_settings
    store_root = os.path.join(django_settings.MEDIA_ROOT, 'jaf_store')

    from django.conf import settings as django_settings
    from datetime import datetime
    media_url = django_settings.MEDIA_URL.rstrip('/')

    divisions = []
    if os.path.isdir(store_root):
        for folder in sorted(os.listdir(store_root)):
            folder_path = os.path.join(store_root, folder)
            if not os.path.isdir(folder_path):
                continue
            files = []
            for fname in sorted(os.listdir(folder_path)):
                if not fname.lower().endswith('.html'):
                    continue
                fpath = os.path.join(folder_path, fname)
                stat = os.stat(fpath)
                files.append({
                    'name': fname,
                    'size_kb': round(stat.st_size / 1024, 1),
                    'stored_at': datetime.fromtimestamp(stat.st_mtime),
                    'url': f'{media_url}/jaf_store/{folder}/{fname}',
                })
            if files:
                divisions.append({'folder': folder, 'files': files})

    return render(request, 'competencies/jaf_store_index.html', {
        'divisions': divisions,
        'store_root': store_root,
    })


@login_required
def jaf_store(request, pk):
    """HRDD stores the JAF as a standalone HTML file in the division folder."""
    if request.method != 'POST':
        return redirect('jaf_detail', pk=pk)

    if request.user.role not in ['hr', 'admin']:
        messages.error(request, 'Only HRDD may store a Job Analysis Form.')
        return redirect('jaf_detail', pk=pk)

    entry = get_object_or_404(JobAnalysisEntry, pk=pk)

    # Determine division folder name
    division = getattr(entry.employee, 'division', None)
    if division:
        folder_name = division.code or division.name
    else:
        folder_name = 'no_division'

    # Sanitise folder / file names
    import re
    safe_folder = re.sub(r'[^\w\-]', '_', folder_name)
    safe_name   = re.sub(r'[^\w\-]', '_', entry.full_name)
    filename    = f'JAF_{entry.pk:04d}_{safe_name}.html'

    # Build destination path
    from django.conf import settings as django_settings
    import os
    store_dir = os.path.join(django_settings.MEDIA_ROOT, 'jaf_store', safe_folder)
    os.makedirs(store_dir, exist_ok=True)
    file_path = os.path.join(store_dir, filename)

    # Render standalone HTML
    from django.template.loader import render_to_string
    html_content = render_to_string('competencies/jaf_stored.html', {
        'entry': entry,
        'secondary_duties': entry.secondary_duties.all(),
        'required_skills': entry.required_skills.select_related('competency').all(),
        'tools': entry.tools_equipment.all(),
        'stored_by': request.user,
        'dotr_logo_b64': _img_b64('img/DOTr.png'),
        'bp_logo_b64': _img_b64('img/BPLogo.png'),
    })

    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(html_content)

    rel_path = os.path.join('jaf_store', safe_folder, filename)
    messages.success(
        request,
        f'JAF stored successfully. File saved to: media/{safe_folder}/{filename}'
    )
    return redirect('jaf_detail', pk=pk)


@login_required
def jaf_export_html(request, pk):
    """Download the JAF as a standalone HTML file."""
    entry = get_object_or_404(JobAnalysisEntry, pk=pk)
    if not _jaf_can_view(request.user, entry):
        messages.error(request, 'Access denied.')
        return redirect('jaf_list')

    from django.template.loader import render_to_string
    html_content = render_to_string('competencies/jaf_stored.html', {
        'entry': entry,
        'secondary_duties': entry.secondary_duties.all(),
        'required_skills': entry.required_skills.select_related('competency').all(),
        'tools': entry.tools_equipment.all(),
        'stored_by': request.user,
        'dotr_logo_b64': _img_b64('img/DOTr.png'),
        'bp_logo_b64': _img_b64('img/BPLogo.png'),
    }, request=request)

    import re
    safe_name = re.sub(r'[^\w\-]', '_', entry.full_name)
    filename = f'JAF_{entry.pk:04d}_{safe_name}.html'

    from django.http import HttpResponse
    response = HttpResponse(html_content, content_type='text/html; charset=utf-8')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@login_required
def jaf_export_pdf(request, pk):
    """Download the JAF as a PDF using xhtml2pdf."""
    entry = get_object_or_404(JobAnalysisEntry, pk=pk)
    if not _jaf_can_view(request.user, entry):
        messages.error(request, 'Access denied.')
        return redirect('jaf_list')

    from django.template.loader import render_to_string
    from django.http import HttpResponse
    import io

    html_content = render_to_string('competencies/jaf_stored.html', {
        'entry': entry,
        'secondary_duties': entry.secondary_duties.all(),
        'required_skills': entry.required_skills.select_related('competency').all(),
        'tools': entry.tools_equipment.all(),
        'stored_by': request.user,
        'dotr_logo_b64': _img_b64('img/DOTr.png'),
        'bp_logo_b64': _img_b64('img/BPLogo.png'),
    }, request=request)

    try:
        from xhtml2pdf import pisa
    except ImportError:
        messages.error(request, 'PDF export requires xhtml2pdf. Run: pip install xhtml2pdf')
        return redirect('jaf_detail', pk=pk)

    buffer = io.BytesIO()
    result = pisa.CreatePDF(html_content, dest=buffer, encoding='utf-8')
    if result.err:
        messages.error(request, 'PDF generation failed. Please try HTML export instead.')
        return redirect('jaf_detail', pk=pk)

    import re
    safe_name = re.sub(r'[^\w\-]', '_', entry.full_name)
    filename = f'JAF_{entry.pk:04d}_{safe_name}.pdf'

    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@login_required
def jaf_export_docx(request, pk):
    """Download the JAF as a Word (.docx) document using python-docx."""
    entry = get_object_or_404(JobAnalysisEntry, pk=pk)
    if not _jaf_can_view(request.user, entry):
        messages.error(request, 'Access denied.')
        return redirect('jaf_list')

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        messages.error(request, 'DOCX export requires python-docx. Run: pip install python-docx')
        return redirect('jaf_detail', pk=pk)

    import io
    import re
    from django.http import HttpResponse

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    BLUE     = RGBColor(0x1e, 0x3a, 0x8a)
    WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
    GRAY_BG  = RGBColor(0xF1, 0xF5, 0xF9)

    def _set_cell_bg(cell, hex_color):
        """Set table cell background colour."""
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tc_pr.append(shd)

    def _add_section_header(doc, title):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = WHITE
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(0)
        # Blue background via paragraph shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1E3A8A')
        pPr.append(shd)
        return p

    def _add_field_row(doc, label, value):
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = 'Table Grid'
        tbl.columns[0].width = Cm(5)
        row = tbl.rows[0]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        _set_cell_bg(row.cells[0], 'F8FAFC')
        row.cells[1].text = value or '—'
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        return tbl

    # ── Title ─────────────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('JOB ANALYSIS FORM')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = BLUE

    org_p = doc.add_paragraph()
    org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_p.add_run('Department of Transportation').font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # ── Header fields ─────────────────────────────────────────────────────────
    _add_field_row(doc, 'Full Name', entry.full_name)
    _add_field_row(doc, 'Position Title', entry.position_title)
    _add_field_row(doc, 'Office / Service / Division', entry.office_service_division or '—')
    _add_field_row(doc, 'Section / Project / Unit', entry.section_project_unit or '—')
    _add_field_row(doc, 'Alternate Position', entry.alternate_position or '—')
    doc.add_paragraph()

    # ── Job Purpose ───────────────────────────────────────────────────────────
    _add_section_header(doc, 'JOB PURPOSE')
    p = doc.add_paragraph(entry.job_purpose or '—')
    p.runs[0].font.size = Pt(9)

    # ── Main Duties ───────────────────────────────────────────────────────────
    _add_section_header(doc, 'MAIN DUTIES AND RESPONSIBILITIES')
    p = doc.add_paragraph(entry.main_duties or '—')
    p.runs[0].font.size = Pt(9)

    # ── Secondary Duties ─────────────────────────────────────────────────────
    _add_section_header(doc, 'SECONDARY DUTIES & RESPONSIBILITIES')
    secondary_duties = entry.secondary_duties.all()
    if secondary_duties:
        headers = ['No.', 'Task', 'Daily', 'Weekly', 'Monthly', 'Quarterly', 'Periodically']
        widths  = [Cm(1), Cm(7), Cm(1.5), Cm(1.5), Cm(1.7), Cm(1.9), Cm(2.5)]
        tbl = doc.add_table(rows=1, cols=7)
        tbl.style = 'Table Grid'
        for i, (h, w) in enumerate(zip(headers, widths)):
            cell = tbl.rows[0].cells[i]
            cell.width = w
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(8)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_bg(cell, 'F1F5F9')
        for idx, duty in enumerate(secondary_duties, start=1):
            row = tbl.add_row()
            row.cells[0].text = str(idx)
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[1].text = duty.task
            freq_map = {'daily': 2, 'weekly': 3, 'monthly': 4, 'quarterly': 5, 'periodically': 6}
            col_idx = freq_map.get(duty.frequency)
            if col_idx is not None:
                row.cells[col_idx].text = '✓'
                row.cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for c in row.cells:
                c.paragraphs[0].runs[0].font.size = Pt(9) if c.paragraphs[0].runs else None
    else:
        doc.add_paragraph('No secondary duties recorded.').runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # ── Required Competencies ─────────────────────────────────────────────────
    _add_section_header(doc, 'REQUIRED COMPETENCIES')
    note_p = doc.add_paragraph('Note: Filled by the immediate supervisor / division chief')
    note_p.runs[0].font.size = Pt(8)
    note_p.runs[0].italic = True
    required_skills = entry.required_skills.select_related('competency').all()
    if required_skills:
        headers = ['No.', 'Category', 'Competency', 'Basic', 'Intermediate', 'Advanced', 'Superior']
        widths  = [Cm(1), Cm(3), Cm(6), Cm(1.5), Cm(2), Cm(2), Cm(1.8)]
        tbl = doc.add_table(rows=1, cols=7)
        tbl.style = 'Table Grid'
        for i, (h, w) in enumerate(zip(headers, widths)):
            cell = tbl.rows[0].cells[i]
            cell.width = w
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(8)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_bg(cell, 'F1F5F9')
        for idx, skill in enumerate(required_skills, start=1):
            row = tbl.add_row()
            row.cells[0].text = str(idx)
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            comp = skill.competency
            row.cells[1].text = comp.get_category_display() if comp else '—'
            row.cells[2].text = (comp.name if comp else skill.skill_name) or '—'
            level_map = {'basic': 3, 'intermediate': 4, 'advanced': 5, 'superior': 6}
            col_idx = level_map.get(skill.proficiency_level)
            if col_idx is not None:
                row.cells[col_idx].text = '✓'
                row.cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for c in row.cells:
                if c.paragraphs[0].runs:
                    c.paragraphs[0].runs[0].font.size = Pt(9)
    else:
        doc.add_paragraph('No competencies recorded.').runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # ── Tools and Equipment ───────────────────────────────────────────────────
    _add_section_header(doc, 'TOOLS AND EQUIPMENT')
    tools = entry.tools_equipment.all()
    if tools:
        for i, t in enumerate(tools, start=1):
            p = doc.add_paragraph(f'{i}. {t.name}', style='List Number')
            if p.runs:
                p.runs[0].font.size = Pt(9)
    else:
        doc.add_paragraph('No tools recorded.').runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # ── Challenges ────────────────────────────────────────────────────────────
    _add_section_header(doc, 'CHALLENGES AND CRITICAL ISSUES')
    p = doc.add_paragraph(entry.challenges_critical_issues or '—')
    if p.runs:
        p.runs[0].font.size = Pt(9)

    # ── Additional Comments ───────────────────────────────────────────────────
    _add_section_header(doc, 'ADDITIONAL COMMENTS')
    p = doc.add_paragraph(entry.additional_comments or '—')
    if p.runs:
        p.runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # ── Certification ─────────────────────────────────────────────────────────
    cert_p = doc.add_paragraph()
    cert_p.add_run(
        'I hereby certify that all information and statements provided herewith are accurate. '
        'I further attest that these details truthfully reflect my current knowledge and understanding.'
    ).font.size = Pt(9)

    cert_tbl = doc.add_table(rows=2, cols=2)
    cert_tbl.style = 'Table Grid'
    cert_tbl.rows[0].cells[0].text = entry.full_name if entry.certified_date else ''
    cert_tbl.rows[0].cells[1].text = entry.certified_date.strftime('%B %d, %Y') if entry.certified_date else ''
    cert_tbl.rows[1].cells[0].text = 'Signature over Printed Name'
    cert_tbl.rows[1].cells[1].text = 'Date'
    for row in cert_tbl.rows:
        for cell in row.cells:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(9)

    # ── Approval block ────────────────────────────────────────────────────────
    if entry.approved_by:
        doc.add_paragraph()
        _add_section_header(doc, 'HRDD APPROVAL')
        appr_tbl = doc.add_table(rows=2, cols=2)
        appr_tbl.style = 'Table Grid'
        appr_tbl.rows[0].cells[0].text = entry.approved_by.get_full_name()
        appr_tbl.rows[0].cells[1].text = entry.approved_date.strftime('%B %d, %Y') if entry.approved_date else ''
        appr_tbl.rows[1].cells[0].text = 'Approved by (HRDD)'
        appr_tbl.rows[1].cells[1].text = 'Date Approved'
        for row in appr_tbl.rows:
            for cell in row.cells:
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.size = Pt(9)

    # ── Serialize ─────────────────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_name = re.sub(r'[^\w\-]', '_', entry.full_name)
    filename = f'JAF_{entry.pk:04d}_{safe_name}.docx'

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def _save_jaf_employee_sections(request, entry):
    """Save secondary duties and tools (employee-editable sections) from POST data."""
    i = 1
    while True:
        task = request.POST.get(f'task_{i}', '').strip()
        freq = request.POST.get(f'freq_{i}', '').strip()
        if not task:
            break
        if freq:
            SecondaryDuty.objects.create(entry=entry, order=i, task=task, frequency=freq)
        i += 1

    i = 1
    while True:
        tool = request.POST.get(f'tool_{i}', '').strip()
        if not tool:
            break
        ToolEquipment.objects.create(entry=entry, order=i, name=tool)
        i += 1


def _save_jaf_skills(request, entry):
    """Save required skills (supervisor-only section) from POST data.
    Reads competency_N (PK from Competency model) + skill_level_N."""
    i = 1
    while True:
        competency_id = request.POST.get(f'competency_{i}', '').strip()
        level = request.POST.get(f'skill_level_{i}', '').strip()
        if not competency_id:
            break
        if level:
            try:
                comp = Competency.objects.get(pk=competency_id, is_active=True)
                RequiredSkill.objects.create(
                    entry=entry, order=i,
                    competency=comp,
                    skill_name=comp.name,
                    proficiency_level=level,
                )
            except Competency.DoesNotExist:
                pass
        i += 1


# ── DOTr Competency Framework Views ──────────────────────────────────────────

_DOTR_CF_ROLES = ['admin', 'hr', 'supervisor', 'executive']


def _supervisor_can_manage_cf(user, competency=None):
    """Return True if a supervisor may create/edit the given (or a new) competency.

    Supervisors may only manage *functional* competencies that belong to their
    own Division.  Core and Leadership competencies are org-wide and restricted
    to admin/HR.
    """
    if user.role != 'supervisor':
        return False
    if not user.division_id:
        return False
    if competency is None:
        return True  # creating a new one is allowed; division is enforced in the view
    return (competency.type == 'functional'
            and competency.division_id == user.division_id)


@login_required
@role_required(_DOTR_CF_ROLES)
def dotr_cf_about(request):
    return render(request, 'competencies/dotr_cf_about.html')


@login_required
@role_required(_DOTR_CF_ROLES)
def dotr_cf_types(request):
    return render(request, 'competencies/dotr_cf_types.html', {
        'core_comps': DOTrCompetency.objects.filter(type='core', is_active=True),
        'leadership_comps': DOTrCompetency.objects.filter(type='leadership', is_active=True),
        'functional_comps': DOTrCompetency.objects.filter(type='functional', is_active=True).order_by('office', 'name'),
    })


@login_required
@role_required(_DOTR_CF_ROLES)
def dotr_cf_levels(request):
    return render(request, 'competencies/dotr_cf_levels.html')


@login_required
@role_required(_DOTR_CF_ROLES)
def dotr_cf_list(request):
    cf_type = request.GET.get('type', '')
    office = request.GET.get('office', '')
    user = request.user
    competencies = DOTrCompetency.objects.filter(is_active=True).prefetch_related('indicators')

    # Supervisors see core/leadership (org-wide) plus only their division's functional competencies.
    if user.role == 'supervisor' and user.division_id:
        competencies = competencies.filter(
            Q(type__in=['core', 'leadership']) |
            Q(type='functional', division_id=user.division_id)
        )

    if cf_type:
        competencies = competencies.filter(type=cf_type)
    if office:
        competencies = competencies.filter(office__icontains=office)

    offices_qs = DOTrCompetency.objects.filter(type='functional', is_active=True)
    if user.role == 'supervisor' and user.division_id:
        offices_qs = offices_qs.filter(division_id=user.division_id)
    offices = offices_qs.values_list('office', flat=True).distinct().order_by('office')

    can_manage = user.role in ('admin', 'hr') or _supervisor_can_manage_cf(user)

    return render(request, 'competencies/dotr_cf_list.html', {
        'competencies': competencies,
        'cf_type': cf_type,
        'office_filter': office,
        'types': DOTrCompetencyType.choices,
        'offices': offices,
        'can_manage': can_manage,
    })


@login_required
@role_required(_DOTR_CF_ROLES)
def dotr_cf_detail(request, pk):
    competency = get_object_or_404(DOTrCompetency, pk=pk)
    user = request.user

    # Supervisors may only view their own division's functional competencies.
    if (user.role == 'supervisor' and user.division_id
            and competency.type == 'functional'
            and competency.division_id != user.division_id):
        messages.error(request, 'You can only view competencies for your own division.')
        return redirect('dotr_cf_list')

    can_manage = user.role in ('admin', 'hr') or _supervisor_can_manage_cf(user, competency)

    indicators_by_level = {}
    for level_val, level_label in DOTrCompetencyIndicator.Level.choices:
        indicators_by_level[level_val] = {
            'label': level_label,
            'items': competency.indicators.filter(level=level_val).order_by('order'),
        }
    return render(request, 'competencies/dotr_cf_detail.html', {
        'competency': competency,
        'indicators_by_level': indicators_by_level,
        'can_manage': can_manage,
    })


@login_required
@role_required(['admin', 'hr', 'supervisor'])
def dotr_cf_create(request):
    user = request.user
    is_supervisor = user.role == 'supervisor'

    if is_supervisor and not user.division_id:
        messages.error(request, 'You must be assigned to a division to create competencies.')
        return redirect('dotr_cf_list')

    if request.method == 'POST':
        cf_type = request.POST.get('type', '')
        # Supervisors are locked to functional type only.
        if is_supervisor:
            cf_type = 'functional'
        name = request.POST.get('name', '').strip()
        description = request.POST.get('description', '').strip()
        order = request.POST.get('order', 0)
        if is_supervisor:
            division = user.division
            office_name = division.name if division else ''
        else:
            division, office_name = _resolve_division(request)
        comp = DOTrCompetency.objects.create(
            name=name, type=cf_type, description=description,
            office=office_name, order=order, created_by=user,
            division=division,
        )
        if division and cf_type == 'functional':
            _save_office_mandates(request, division)
        _save_dotr_indicators(request, comp)
        messages.success(request, f'Competency "{name}" created successfully.')
        return redirect('dotr_cf_detail', pk=comp.pk)

    locked_division = user.division if is_supervisor else None
    existing_mandates = list(locked_division.mandates.order_by('order')) if locked_division else []
    type_choices = [('functional', 'Functional Competency')] if is_supervisor else list(DOTrCompetencyType.choices)
    return render(request, 'competencies/dotr_cf_form.html', _cf_form_context(
        title='Add Functional Competency',
        types=type_choices,
        level_choices=DOTrCompetencyIndicator.Level.choices,
        is_supervisor=is_supervisor,
        locked_division=locked_division,
        existing_mandates=existing_mandates,
    ))


@login_required
@role_required(['admin', 'hr', 'supervisor'])
def dotr_cf_edit(request, pk):
    competency = get_object_or_404(DOTrCompetency, pk=pk)
    user = request.user
    is_supervisor = user.role == 'supervisor'

    # Supervisors may only edit functional competencies belonging to their division.
    if is_supervisor and not _supervisor_can_manage_cf(user, competency):
        messages.error(request, 'You can only edit functional competencies for your own division.')
        return redirect('dotr_cf_detail', pk=pk)

    if request.method == 'POST':
        cf_type = request.POST.get('type', '')
        if is_supervisor:
            cf_type = 'functional'
        competency.name = request.POST.get('name', '').strip()
        competency.type = cf_type
        competency.description = request.POST.get('description', '').strip()
        competency.order = request.POST.get('order', 0)
        if is_supervisor:
            division = user.division
            office_name = division.name if division else ''
        else:
            division, office_name = _resolve_division(request)
        competency.division = division if cf_type == 'functional' else None
        competency.office = office_name if cf_type == 'functional' else ''
        if division and cf_type == 'functional':
            _save_office_mandates(request, division)
        competency.save()
        competency.indicators.all().delete()
        _save_dotr_indicators(request, competency)
        messages.success(request, f'Competency "{competency.name}" updated.')
        return redirect('dotr_cf_detail', pk=competency.pk)

    indicators_by_level = {}
    for level_val, level_label in DOTrCompetencyIndicator.Level.choices:
        indicators_by_level[level_val] = {
            'label': level_label,
            'items': list(competency.indicators.filter(level=level_val).order_by('order')),
        }
    existing_mandates = list(competency.division.mandates.order_by('order')) if competency.division else []
    locked_division = user.division if is_supervisor else None
    type_choices = [('functional', 'Functional Competency')] if is_supervisor else list(DOTrCompetencyType.choices)
    return render(request, 'competencies/dotr_cf_form.html', _cf_form_context(
        title='Edit Competency',
        types=type_choices,
        level_choices=DOTrCompetencyIndicator.Level.choices,
        competency=competency,
        indicators_by_level=indicators_by_level,
        existing_mandates=existing_mandates,
        is_supervisor=is_supervisor,
        locked_division=locked_division,
    ))


@login_required
@role_required(['admin', 'hr'])
def dotr_cf_delete(request, pk):
    competency = get_object_or_404(DOTrCompetency, pk=pk)
    if request.method == 'POST':
        name = competency.name
        competency.delete()
        messages.success(request, f'Competency "{name}" deleted.')
        return redirect('dotr_cf_list')
    return render(request, 'competencies/dotr_cf_confirm_delete.html', {'competency': competency})


@login_required
@role_required(_DOTR_CF_ROLES)
def dotr_cf_office(request):
    user = request.user
    offices_qs = (
        Division.objects
        .filter(dotr_competencies__type='functional', dotr_competencies__is_active=True)
        .distinct()
        .order_by('name')
    )
    # Supervisors only see their own division in the office selector.
    if user.role == 'supervisor' and user.division_id:
        offices_qs = offices_qs.filter(pk=user.division_id)

    selected_office_name = request.GET.get('office', '')
    # Auto-select a supervisor's own division when they first arrive.
    if user.role == 'supervisor' and user.division_id and not selected_office_name:
        try:
            selected_office_name = user.division.name
        except Exception:
            pass

    competencies_with_levels = []
    mandates = []
    division = None
    if selected_office_name:
        try:
            division = Division.objects.prefetch_related('mandates').get(name=selected_office_name)
            # Prevent supervisors from viewing another division's page via URL manipulation.
            if user.role == 'supervisor' and user.division_id and division.pk != user.division_id:
                messages.error(request, 'You can only view your own division\'s competencies.')
                return redirect('dotr_cf_office')
            mandates = list(division.mandates.order_by('order'))
        except Division.DoesNotExist:
            pass
        qs = DOTrCompetency.objects.filter(
            type='functional', is_active=True, office=selected_office_name
        ).prefetch_related('indicators').order_by('order', 'name')
        for comp in qs:
            levels = {}
            for level_val, level_label in DOTrCompetencyIndicator.Level.choices:
                levels[level_val] = {
                    'label': level_label,
                    'items': list(comp.indicators.filter(level=level_val).order_by('order')),
                }
            competencies_with_levels.append({'comp': comp, 'levels': levels})

    can_manage = user.role in ('admin', 'hr') or _supervisor_can_manage_cf(user)

    return render(request, 'competencies/dotr_cf_office.html', {
        'offices': offices_qs,
        'selected_office': selected_office_name,
        'division': division,
        'mandates': mandates,
        'competencies_with_levels': competencies_with_levels,
        'level_choices': DOTrCompetencyIndicator.Level.choices,
        'can_manage': can_manage,
    })


def _resolve_division(request):
    """Return (Division | None, division_name str) from POST office_unit_id field."""
    uid = request.POST.get('office_unit_id', '').strip()
    if uid:
        try:
            div = Division.objects.get(pk=int(uid))
            return div, div.name
        except (Division.DoesNotExist, ValueError):
            pass
    return None, ''


def _cf_form_context(**kwargs):
    """Build the shared template context for dotr_cf_form, including divisions + JSON mandates."""
    import json as _json
    divisions = list(Division.objects.prefetch_related('mandates').order_by('name'))
    mandates_json = {
        str(d.pk): [m.description for m in d.mandates.all()]
        for d in divisions
    }
    ctx = {
        'existing_offices': divisions,
        'offices_mandates_json': _json.dumps(mandates_json),
        'existing_mandates': [],
    }
    ctx.update(kwargs)
    return ctx


def _save_office_mandates(request, division):
    """Replace all mandates for the given division from POST mandate_N fields."""
    division.mandates.all().delete()
    i = 1
    while True:
        desc = request.POST.get(f'mandate_{i}', '').strip()
        if not desc:
            break
        DOTrOfficeMandate.objects.create(division=division, order=i, description=desc)
        i += 1


def _save_dotr_indicators(request, competency):
    """Parse indicator rows from POST and create DOTrCompetencyIndicator records."""
    for level_val, _ in DOTrCompetencyIndicator.Level.choices:
        i = 1
        while True:
            desc = request.POST.get(f'indicator_{level_val}_{i}', '').strip()
            num = request.POST.get(f'indicator_num_{level_val}_{i}', '').strip()
            if not desc:
                break
            DOTrCompetencyIndicator.objects.create(
                competency=competency,
                level=level_val,
                indicator_number=num,
                description=desc,
                order=i,
            )
            i += 1
